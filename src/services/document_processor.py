# Enhanced Multi-Modal Document Processor Service
"""
SCRUM-40: Enhanced Multi-Modal Document Processing and Text Extraction

Service for comprehensive document processing with full multi-modal support:
1. Extract text from PDF, DOCX, and TXT files
2. Extract and analyze images, charts, and visual elements
3. Extract and structure tables with data preservation
4. Generate AI descriptions for visual elements
5. Create Warren-optimized context summaries
6. Validate file types and security

Created: July 3, 2025
Status: Phase 1 - Core Structure Implementation
"""

import logging
import asyncio
import json
import io
import re
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

# Multi-modal document processing imports
try:
    import fitz  # PyMuPDF for PDF processing with image extraction
    import docx  # python-docx for Word document processing
    from PIL import Image, ImageDraw, ImageFont  # Pillow for image processing
    import tabulate  # Table formatting and extraction
    import magic  # File type detection
    import chardet  # Character encoding detection
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    logging.warning(f"Document processing dependencies not available: {e}")
    DEPENDENCIES_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Enhanced multi-modal document processing with visual element extraction.
    
    Capabilities:
    - PDF: Text + images + charts + tables extraction
    - DOCX: Text + embedded content + tables + images
    - TXT: Enhanced text processing with metadata
    - Visual Analysis: Chart/graph detection and description
    - Table Processing: Structured data extraction and formatting
    - Security: File validation and type checking
    """
    
    def __init__(self):
        """Initialize the document processor with multi-modal capabilities."""
        self.supported_types = ['pdf', 'docx', 'txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.image_description_prompt = """
        Analyze this image from a financial document and provide a concise description 
        focusing on charts, graphs, tables, or financial data visualizations.
        """
        
        if not DEPENDENCIES_AVAILABLE:
            logger.error("Document processing dependencies not installed")
            raise ImportError("Required dependencies for document processing not available")
    
    # ===== CORE PROCESSING METHODS =====
    
    async def process_uploaded_file(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Main entry point for processing any uploaded file.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: File type ('pdf', 'docx', 'txt')
            
        Returns:
            Dict with comprehensive multi-modal extraction results
        """
        try:
            logger.info(f"Processing {content_type} file: {filename}")
            
            # Validate file
            if not self.validate_file_type(filename, file_content):
                raise ValueError(f"Invalid file type or security check failed: {filename}")
            
            # Route to appropriate processor
            if content_type.lower() == 'pdf':
                result = await self.extract_content_from_pdf(file_content)
            elif content_type.lower() == 'docx':
                result = await self.extract_content_from_docx(file_content)
            elif content_type.lower() == 'txt':
                result = await self.process_text_file(file_content)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")
            
            # Add processing metadata
            result['metadata']['original_filename'] = filename
            result['metadata']['content_type'] = content_type
            result['metadata']['processed_at'] = datetime.utcnow().isoformat()
            result['metadata']['processor_version'] = "1.0.0"
            
            logger.info(f"Successfully processed {filename}: {result['metadata']['word_count']} words, "
                       f"{result['metadata']['total_images']} images, {result['metadata']['total_tables']} tables")
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
    
    async def extract_content_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract comprehensive content from PDF with multi-modal analysis.
        
        Returns:
            Dict containing text, images, tables, and visual summaries
        """
        logger.info("Starting PDF multi-modal extraction")
        start_time = datetime.utcnow()
        
        # Initialize result structure
        result = {
            "text": "",
            "images": [],
            "tables": [],
            "visual_summary": "",
            "warren_context": "",
            "metadata": {
                "word_count": 0,
                "file_size_bytes": len(file_content),
                "total_images": 0,
                "total_tables": 0,
                "total_pages": 0,
                "processing_time_ms": 0
            }
        }
        
        try:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            result['metadata']['total_pages'] = len(pdf_document)
            
            full_text = ""
            all_images = []
            all_tables = []
            
            # Process each page
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Extract text from page
                page_text = page.get_text()
                full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Extract images from page
                page_images = await self._extract_images_from_pdf_page(page, page_num + 1)
                all_images.extend(page_images)
                
                # Extract tables from page
                page_tables = await self._extract_tables_from_pdf_page(page, page_num + 1)
                all_tables.extend(page_tables)
            
            pdf_document.close()
            
            # Process extracted content
            result['text'] = self._clean_extracted_text(full_text)
            result['images'] = all_images
            result['tables'] = all_tables
            result['metadata']['word_count'] = len(result['text'].split())
            result['metadata']['total_images'] = len(all_images)
            result['metadata']['total_tables'] = len(all_tables)
            
            # Generate summaries
            result['visual_summary'] = await self._create_visual_summary(all_images, all_tables)
            result['warren_context'] = await self._create_warren_context_summary(result)
            
            # Calculate processing time
            end_time = datetime.utcnow()
            result['metadata']['processing_time_ms'] = int((end_time - start_time).total_seconds() * 1000)
            
            logger.info(f"PDF extraction complete: {result['metadata']['total_pages']} pages, "
                       f"{result['metadata']['word_count']} words, {result['metadata']['total_images']} images")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            # Return partial results if possible
            result['text'] = "Error: Could not extract PDF content"
            result['metadata']['processing_error'] = str(e)
            return result
    
    async def extract_content_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract comprehensive content from DOCX with embedded content analysis.
        
        Returns:
            Dict containing text, images, tables, and visual summaries
        """
        logger.info("Starting DOCX multi-modal extraction")
        start_time = datetime.utcnow()
        
        # Initialize result structure
        result = {
            "text": "",
            "images": [],
            "tables": [],
            "visual_summary": "",
            "warren_context": "",
            "metadata": {
                "word_count": 0,
                "file_size_bytes": len(file_content),
                "total_images": 0,
                "total_tables": 0,
                "total_paragraphs": 0,
                "processing_time_ms": 0
            }
        }
        
        try:
            # Open DOCX document
            doc_stream = io.BytesIO(file_content)
            document = docx.Document(doc_stream)
            
            full_text = ""
            all_images = []
            all_tables = []
            
            # Extract text from paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            result['metadata']['total_paragraphs'] = len(document.paragraphs)
            
            # Extract tables
            for table_idx, table in enumerate(document.tables):
                table_data = await self._extract_docx_table(table, table_idx)
                if table_data:
                    all_tables.append(table_data)
            
            # Extract embedded images/shapes
            embedded_images = await self._extract_docx_images(document)
            all_images.extend(embedded_images)
            
            # Process extracted content
            result['text'] = self._clean_extracted_text(full_text)
            result['images'] = all_images
            result['tables'] = all_tables
            result['metadata']['word_count'] = len(result['text'].split())
            result['metadata']['total_images'] = len(all_images)
            result['metadata']['total_tables'] = len(all_tables)
            
            # Generate summaries
            result['visual_summary'] = await self._create_visual_summary(all_images, all_tables)
            result['warren_context'] = await self._create_warren_context_summary(result)
            
            # Calculate processing time
            end_time = datetime.utcnow()
            result['metadata']['processing_time_ms'] = int((end_time - start_time).total_seconds() * 1000)
            
            logger.info(f"DOCX extraction complete: {result['metadata']['total_paragraphs']} paragraphs, "
                       f"{result['metadata']['word_count']} words, {result['metadata']['total_images']} images")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            # Return partial results if possible
            result['text'] = "Error: Could not extract DOCX content"
            result['metadata']['processing_error'] = str(e)
            return result
    
    async def process_text_file(self, file_content: bytes) -> Dict[str, Any]:
        """
        Process text file with enhanced metadata extraction.
        
        Returns:
            Dict containing text and enhanced metadata
        """
        logger.info("Starting TXT file processing with enhanced metadata")
        start_time = datetime.utcnow()
        
        # Initialize result structure
        result = {
            "text": "",
            "images": [],  # Empty for text files
            "tables": [],  # Will detect table-like structures
            "visual_summary": "",
            "warren_context": "",
            "metadata": {
                "word_count": 0,
                "file_size_bytes": len(file_content),
                "total_images": 0,
                "total_tables": 0,
                "line_count": 0,
                "character_encoding": "unknown",
                "processing_time_ms": 0
            }
        }
        
        try:
            # Detect character encoding
            encoding_result = chardet.detect(file_content)
            encoding = encoding_result.get('encoding', 'utf-8')
            result['metadata']['character_encoding'] = encoding
            
            # Decode text content
            text_content = file_content.decode(encoding, errors='ignore')
            
            # Clean and process text
            cleaned_text = self._clean_extracted_text(text_content)
            result['text'] = cleaned_text
            
            # Extract metadata
            lines = text_content.split('\n')
            result['metadata']['line_count'] = len(lines)
            result['metadata']['word_count'] = len(cleaned_text.split())
            
            # Detect table-like structures in text
            text_tables = await self._detect_text_tables(text_content)
            result['tables'] = text_tables
            result['metadata']['total_tables'] = len(text_tables)
            
            # Generate summaries
            result['visual_summary'] = await self._create_visual_summary([], text_tables)
            result['warren_context'] = await self._create_warren_context_summary(result)
            
            # Calculate processing time
            end_time = datetime.utcnow()
            result['metadata']['processing_time_ms'] = int((end_time - start_time).total_seconds() * 1000)
            
            logger.info(f"TXT processing complete: {result['metadata']['line_count']} lines, "
                       f"{result['metadata']['word_count']} words, encoding: {encoding}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            # Return partial results if possible
            result['text'] = "Error: Could not process text file"
            result['metadata']['processing_error'] = str(e)
            return result

    # ===== VISUAL ELEMENT PROCESSING METHODS =====
    
    async def _extract_images_from_pdf_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract and analyze images from a PDF page."""
        images = []
        
        try:
            # Get image list from page
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Extract image
                    xref = img[0]
                    pix = fitz.Pixmap(page.document, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Convert to PIL Image for analysis
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                        
                        # Analyze image
                        image_info = {
                            "type": self._classify_image_type(pil_image),
                            "description": await self._generate_image_description(pil_image),
                            "position": f"page_{page_num}_img_{img_index + 1}",
                            "metadata": {
                                "width": pix.width,
                                "height": pix.height,
                                "colorspace": pix.colorspace.name if pix.colorspace else "unknown",
                                "size_bytes": len(img_data)
                            }
                        }
                        
                        # Try to extract data if it's a chart/graph
                        if image_info["type"] in ["chart", "graph", "table"]:
                            chart_data = await self._extract_chart_data(pil_image)
                            if chart_data:
                                image_info["data_extracted"] = chart_data
                        
                        images.append(image_info)
                    
                    pix = None  # Release memory
                    
                except Exception as e:
                    logger.warning(f"Could not process image {img_index} on page {page_num}: {str(e)}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Could not extract images from page {page_num}: {str(e)}")
        
        return images
    
    async def _extract_tables_from_pdf_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract and structure tables from a PDF page."""
        tables = []
        
        try:
            # Use PyMuPDF's table detection
            table_list = page.find_tables()
            
            for table_index, table in enumerate(table_list):
                try:
                    # Extract table data
                    table_data = table.extract()
                    
                    if table_data and len(table_data) > 1:  # Must have header + data
                        table_info = {
                            "data": table_data,
                            "description": await self._generate_table_description(table_data),
                            "position": f"page_{page_num}_table_{table_index + 1}",
                            "metadata": {
                                "rows": len(table_data),
                                "columns": len(table_data[0]) if table_data else 0,
                                "bbox": table.bbox
                            }
                        }
                        
                        tables.append(table_info)
                        
                except Exception as e:
                    logger.warning(f"Could not process table {table_index} on page {page_num}: {str(e)}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Could not extract tables from page {page_num}: {str(e)}")
        
        return tables
    
    async def _extract_docx_table(self, table, table_idx: int) -> Optional[Dict[str, Any]]:
        """Extract table data from DOCX document."""
        try:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data and len(table_data) > 1:  # Must have header + data
                return {
                    "data": table_data,
                    "description": await self._generate_table_description(table_data),
                    "position": f"docx_table_{table_idx + 1}",
                    "metadata": {
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0
                    }
                }
                
        except Exception as e:
            logger.warning(f"Could not extract DOCX table {table_idx}: {str(e)}")
        
        return None
    
    async def _extract_docx_images(self, document) -> List[Dict[str, Any]]:
        """Extract embedded images from DOCX document."""
        images = []
        
        try:
            # Extract images from document relationships
            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        pil_image = Image.open(io.BytesIO(image_data))
                        
                        image_info = {
                            "type": self._classify_image_type(pil_image),
                            "description": await self._generate_image_description(pil_image),
                            "position": f"docx_embedded_{len(images) + 1}",
                            "metadata": {
                                "width": pil_image.width,
                                "height": pil_image.height,
                                "format": pil_image.format,
                                "size_bytes": len(image_data)
                            }
                        }
                        
                        images.append(image_info)
                        
                    except Exception as e:
                        logger.warning(f"Could not process embedded image: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.warning(f"Could not extract DOCX images: {str(e)}")
        
        return images
    
    async def _detect_text_tables(self, text_content: str) -> List[Dict[str, Any]]:
        """Detect table-like structures in plain text."""
        tables = []
        lines = text_content.split('\n')
        
        try:
            current_table = []
            in_table = False
            
            for line_num, line in enumerate(lines):
                # Detect table patterns (tabs, multiple spaces, pipe separators)
                if self._is_table_line(line):
                    current_table.append(line.split())
                    in_table = True
                else:
                    if in_table and len(current_table) > 1:
                        # End of table, process it
                        table_info = {
                            "data": current_table,
                            "description": await self._generate_table_description(current_table),
                            "position": f"text_table_{len(tables) + 1}",
                            "metadata": {
                                "rows": len(current_table),
                                "columns": len(current_table[0]) if current_table else 0,
                                "start_line": line_num - len(current_table) + 1,
                                "end_line": line_num
                            }
                        }
                        tables.append(table_info)
                    
                    current_table = []
                    in_table = False
            
            # Handle table at end of file
            if in_table and len(current_table) > 1:
                table_info = {
                    "data": current_table,
                    "description": await self._generate_table_description(current_table),
                    "position": f"text_table_{len(tables) + 1}",
                    "metadata": {
                        "rows": len(current_table),
                        "columns": len(current_table[0]) if current_table else 0
                    }
                }
                tables.append(table_info)
                
        except Exception as e:
            logger.warning(f"Could not detect text tables: {str(e)}")
        
        return tables
    
    # ===== ANALYSIS AND DESCRIPTION METHODS =====
    
    def _classify_image_type(self, image: "Image.Image") -> str:
        """Classify image type based on visual characteristics."""
        # Simple heuristic-based classification
        # In production, this could use ML models
        
        width, height = image.size
        aspect_ratio = width / height if height > 0 else 1
        
        # Basic classification based on size and aspect ratio
        if width > 400 and height > 300:
            if 1.2 <= aspect_ratio <= 2.0:
                return "chart"  # Likely a chart or graph
            elif aspect_ratio > 2.0:
                return "diagram"  # Wide layout, likely a diagram
            else:
                return "image"  # More square, likely a photo or icon
        else:
            return "icon"  # Small image, likely an icon
    
    async def _generate_image_description(self, image: "Image.Image") -> str:
        """Generate AI-powered description of the image."""
        # For now, return a placeholder description
        # In production, this would use Claude Vision API or similar
        
        image_type = self._classify_image_type(image)
        width, height = image.size
        
        descriptions = {
            "chart": f"Financial chart or graph ({width}x{height}px) - likely contains performance data, trends, or comparisons",
            "graph": f"Data visualization graph ({width}x{height}px) - may show metrics, analytics, or statistical information", 
            "diagram": f"Process or organizational diagram ({width}x{height}px) - could illustrate workflows, hierarchies, or concepts",
            "table": f"Visual table or data grid ({width}x{height}px) - structured information in tabular format",
            "image": f"Financial document image ({width}x{height}px) - may contain screenshots, photos, or illustrations",
            "icon": f"Small icon or symbol ({width}x{height}px) - likely decorative or navigational element"
        }
        
        return descriptions.get(image_type, f"Visual element ({width}x{height}px) - content analysis needed")
    
    async def _generate_table_description(self, table_data: List[List[str]]) -> str:
        """Generate description of table content and purpose."""
        if not table_data or len(table_data) < 2:
            return "Empty or invalid table"
        
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        
        # Analyze headers to understand table purpose
        headers = table_data[0] if table_data else []
        header_text = " ".join(headers).lower()
        
        # Financial keywords detection
        financial_keywords = ['revenue', 'profit', 'cost', 'price', 'rate', 'return', 'growth', 
                             'quarter', 'year', 'performance', 'portfolio', 'investment']
        
        purpose = "data table"
        if any(keyword in header_text for keyword in financial_keywords):
            purpose = "financial data table"
        
        return f"{purpose.title()} with {rows} rows and {cols} columns - headers: {', '.join(headers[:3])}{'...' if len(headers) > 3 else ''}"
    
    async def _extract_chart_data(self, image: "Image.Image") -> Optional[str]:
        """Extract data from charts/graphs (placeholder for future ML implementation)."""
        # Placeholder for future OCR/ML-based chart data extraction
        # This would use specialized libraries like plotdigitizer or ML models
        
        return "Chart data extraction not yet implemented - visual analysis available"
    
    # ===== SUMMARY GENERATION METHODS =====
    
    async def _create_visual_summary(self, images: List[Dict], tables: List[Dict]) -> str:
        """Create human-readable summary of visual elements."""
        if not images and not tables:
            return "No visual elements detected"
        
        summary_parts = []
        
        if images:
            image_types = {}
            for img in images:
                img_type = img.get('type', 'unknown')
                image_types[img_type] = image_types.get(img_type, 0) + 1
            
            image_summary = ", ".join([f"{count} {type}{'s' if count > 1 else ''}" 
                                     for type, count in image_types.items()])
            summary_parts.append(f"Images: {image_summary}")
        
        if tables:
            total_rows = sum(table.get('metadata', {}).get('rows', 0) for table in tables)
            summary_parts.append(f"Tables: {len(tables)} table{'s' if len(tables) > 1 else ''} with {total_rows} total rows")
        
        return "Document contains " + " and ".join(summary_parts)
    
    async def _create_warren_context_summary(self, extracted_data: Dict) -> str:
        """Create Warren-optimized context summary including visual awareness."""
        context_parts = []
        
        # Text summary
        word_count = extracted_data.get('metadata', {}).get('word_count', 0)
        if word_count > 0:
            context_parts.append(f"Text content: {word_count} words")
        
        # Visual elements summary
        images = extracted_data.get('images', [])
        tables = extracted_data.get('tables', [])
        
        if images:
            chart_count = sum(1 for img in images if img.get('type') in ['chart', 'graph'])
            if chart_count > 0:
                context_parts.append(f"Visual data: {chart_count} chart{'s' if chart_count > 1 else ''}/graph{'s' if chart_count > 1 else ''}")
        
        if tables:
            data_rows = sum(table.get('metadata', {}).get('rows', 0) for table in tables)
            context_parts.append(f"Structured data: {len(tables)} table{'s' if len(tables) > 1 else ''} with {data_rows} data points")
        
        # Create Warren context
        if context_parts:
            base_context = "Financial document with " + ", ".join(context_parts)
            
            # Add specific visual references for Warren
            visual_refs = []
            for img in images:
                if img.get('description'):
                    visual_refs.append(f"'{img['description']}'")
            
            for table in tables:
                if table.get('description'):
                    visual_refs.append(f"'{table['description']}'")
            
            if visual_refs:
                base_context += f". Contains: {', '.join(visual_refs[:3])}"
                if len(visual_refs) > 3:
                    base_context += f" and {len(visual_refs) - 3} more visual elements"
            
            return base_context
        
        return "Document processed - content available for analysis"
    
    # ===== UTILITY AND VALIDATION METHODS =====
    
    def _clean_extracted_text(self, raw_text: str) -> str:
        """Clean and normalize extracted text."""
        if not raw_text:
            return ""
        
        # Remove excessive whitespace
        cleaned = re.sub(r'\s+', ' ', raw_text)
        
        # Remove page markers and artifacts
        cleaned = re.sub(r'--- Page \d+ ---', '', cleaned)
        
        # Normalize line breaks
        cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)
        
        return cleaned.strip()
    
    def _is_table_line(self, line: str) -> bool:
        """Detect if a line contains table-like structure."""
        if not line.strip():
            return False
        
        # Check for common table separators
        tab_count = line.count('\t')
        pipe_count = line.count('|')
        space_groups = len(re.findall(r'\s{2,}', line))
        
        # Heuristic: likely a table if multiple separators
        return tab_count >= 2 or pipe_count >= 2 or space_groups >= 2
    
    def validate_file_type(self, filename: str, content: bytes) -> bool:
        """
        Validate file type and perform security checks.
        
        Args:
            filename: Original filename
            content: File content bytes
            
        Returns:
            bool: True if file is valid and safe
        """
        try:
            # Check file size
            if len(content) > self.max_file_size:
                logger.error(f"File too large: {len(content)} bytes > {self.max_file_size}")
                return False
            
            # Check filename extension
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            if file_ext not in self.supported_types:
                logger.error(f"Unsupported file extension: {file_ext}")
                return False
            
            # Use python-magic for MIME type detection (if available)
            try:
                detected_type = magic.from_buffer(content[:1024], mime=True)
                
                # Map MIME types to extensions
                mime_mappings = {
                    'application/pdf': 'pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'text/plain': 'txt'
                }
                
                detected_ext = None
                for mime_type, ext in mime_mappings.items():
                    if detected_type.startswith(mime_type):
                        detected_ext = ext
                        break
                
                if detected_ext and detected_ext != file_ext:
                    logger.warning(f"File extension mismatch: {file_ext} vs detected {detected_ext}")
                    # Allow it but log the discrepancy
                
            except Exception as e:
                logger.warning(f"Could not detect MIME type: {e}")
                # Continue with filename-based validation
            
            # Basic content validation
            if len(content) < 10:  # Too small to be a real document
                logger.error("File content too small to be valid")
                return False
            
            # PDF-specific validation
            if file_ext == 'pdf':
                if not content.startswith(b'%PDF-'):
                    logger.error("Invalid PDF header")
                    return False
            
            # DOCX-specific validation (ZIP-based format)
            elif file_ext == 'docx':
                if not content.startswith(b'PK'):
                    logger.error("Invalid DOCX header (not ZIP format)")
                    return False
            
            logger.info(f"File validation successful: {filename} ({len(content)} bytes)")
            return True
            
        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False
